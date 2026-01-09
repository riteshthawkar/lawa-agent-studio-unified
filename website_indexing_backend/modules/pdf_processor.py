"""
PDF and document processing module.
Handles PDFs, Word documents, Excel files, and other document types.
Uses OCR and AI for text extraction and conversion to Markdown.
"""

import os
import tempfile
import sys
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from concurrent.futures import ThreadPoolExecutor, as_completed
import logging
import time

import fitz  # PyMuPDF
import google.generativeai as genai
import pandas as pd
from PIL import Image

from .config import ProcessingConfig

logger = logging.getLogger(__name__)


# --- Helper Functions (Picklable for Multiprocessing) ---

def setup_gemini_client(api_key: Optional[str]) -> Optional[genai.GenerativeModel]:
    """Setup Gemini AI client."""
    try:
        if api_key:
            genai.configure(api_key=api_key)
            return genai.GenerativeModel('models/gemini-2.0-flash')
        return None
    except Exception as e:
        logger.error(f"Error setting up Gemini AI: {e}")
        return None

def ocr_image_to_markdown(image_path: str, gemini_model: Optional[genai.GenerativeModel], is_embedded_graphic: bool = False) -> str:
    """Perform OCR on an image using Gemini AI."""
    if not gemini_model:
        return "[OCR Error: Gemini AI not configured]"

    try:
        prompt_text = (
            "This image is an embedded graphic (e.g., chart, diagram, illustration, graph) from a document page. "
            "IMPORTANT: Perform a detailed analysis of this graphic and extract ALL information: "
            "1. Extract ALL text visible in the graphic (labels, titles, legends, annotations, data points). "
            "2. For charts and graphs: Describe the type of chart, analyze key trends, describe the axes, and extract numerical values when possible. "
            "3. For diagrams and flowcharts: Describe the structure, connections, and relationships between elements. "
            "4. For tables: Format as a markdown table with proper column headers and row values. "
            "5. For any other data visualization: Extract the key insights, patterns, and quantitative information. "
            "6. If there are color-coded elements, mention their significance. "
            "7. Include any footnotes or source attributions present in the graphic. "
            "Present this information in clear, structured markdown format. "
            "If truly no information can be extracted, state that briefly (e.g., '[Decorative image with no extractable content]'). "
            "Do NOT just return an image tag or merely describe what the image visually looks like without extracting the data."
        ) if is_embedded_graphic else (
            "Accurately convert this document image to markdown. "
            "Use appropriate markdown syntax for headings, lists, tables, and formatting. "
            "IMPORTANT: Pay special attention to any charts, graphs, diagrams, or other graphical elements: "
            "1. Extract ALL text from these elements including labels, legends, and annotations. "
            "2. For data visualizations: describe the type of visualization, extract key data points, trends, and patterns. "
            "3. Convert tables to proper markdown tables with headers and data. "
            "4. For complex graphics: summarize key information and insights conveyed by the graphic. "
            "5. Preserve the document's structure, hierarchy, and relationships between text and graphics. "
            "Preserve original layout and organization. "
            "Do NOT wrap the content in markdown code blocks or use triple backticks."
        )

        image_part = genai.upload_file(str(image_path))
        response = gemini_model.generate_content([prompt_text, image_part])

        if response and hasattr(response, 'text') and response.text:
            cleaned_text = response.text.strip()
            
            # Remove code block wrappers
            if cleaned_text.startswith("```markdown"):
                cleaned_text = cleaned_text[11:].strip()
            elif cleaned_text.startswith("```"):
                cleaned_text = cleaned_text[3:].strip()
            if cleaned_text.endswith("```"):
                cleaned_text = cleaned_text[:-3].strip()

            return cleaned_text
        else:
            return "[OCR Error: Empty response from model]"

    except Exception as e:
        logger.error(f"OCR error for {image_path}: {e}")
        return f"[OCR Error: {e}]"
    finally:
        # Clean up uploaded file
        try:
            if 'image_part' in locals() and hasattr(image_part, 'name'):
                genai.delete_file(image_part.name)
        except Exception:
            pass

def check_memory_usage(threshold_mb: int = 500) -> bool:
    """Check if current memory usage is within acceptable limits."""
    try:
        import psutil
        process = psutil.Process()
        memory_mb = process.memory_info().rss / 1024 / 1024
        if memory_mb > threshold_mb:
            logger.warning(f"High memory usage during PDF processing: {memory_mb:.1f}MB")
            return False
        return True
    except ImportError:
        return True
    except Exception:
        return True

def cleanup_memory():
    """Force garbage collection."""
    import gc
    gc.collect()

def has_significant_images(doc, image_list: list, min_dimension: int = 200) -> bool:
    """Check if page has significant images."""
    if not image_list:
        return False

    for img_info in image_list:
        try:
            xref = img_info[0]
            base_image = doc.extract_image(xref)
            if base_image:
                width = base_image.get("width", 0)
                height = base_image.get("height", 0)
                if width >= min_dimension or height >= min_dimension:
                    return True
        except Exception:
            continue
    return False

def extract_and_ocr_embedded_images_memory(
    doc, page_num: int, source_url: str,
    gemini_model: Optional[genai.GenerativeModel],
    min_size: int = 100,
    max_images_per_page: int = 5,
    prioritize_large: bool = True,
    total_ocr_count: int = 0,
    max_total_images: int = 20
) -> Tuple[List[str], int]:
    """Extract embedded images and OCR them."""
    ocr_results = []
    images_ocrd_this_page = 0

    try:
        page = doc.load_page(page_num)
        image_list = page.get_images(full=True)

        if not image_list:
            return [], total_ocr_count

        logger.debug(f"Found {len(image_list)} embedded images on page {page_num + 1}")

        image_candidates = []
        for img_idx, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue

                width = base_image.get("width", 0)
                height = base_image.get("height", 0)

                # Skip small images
                if width < min_size or height < min_size:
                    continue

                area = width * height
                image_candidates.append({
                    "idx": img_idx,
                    "xref": xref,
                    "width": width,
                    "height": height,
                    "area": area,
                    "image_bytes": base_image["image"],
                    "ext": base_image.get("ext", "png")
                })
            except Exception as e:
                logger.debug(f"Error extracting image metadata for xref {xref}: {e}")
                continue

        if not image_candidates:
            return [], total_ocr_count

        if prioritize_large:
            image_candidates.sort(key=lambda x: x["area"], reverse=True)

        for img_data in image_candidates:
            if images_ocrd_this_page >= max_images_per_page:
                break
            if total_ocr_count >= max_total_images:
                break

            try:
                with tempfile.NamedTemporaryFile(suffix=f".{img_data['ext']}", delete=False) as tmp_img:
                    tmp_img.write(img_data["image_bytes"])
                    tmp_img_path = tmp_img.name

                try:
                    ocr_content = ocr_image_to_markdown(tmp_img_path, gemini_model, is_embedded_graphic=True)
                    if ocr_content and len(ocr_content) > 20:
                        ocr_results.append(
                            f"\n**Figure {img_data['idx'] + 1}** ({img_data['width']}x{img_data['height']}):\n"
                            f"{ocr_content}\n"
                        )
                        images_ocrd_this_page += 1
                        total_ocr_count += 1
                finally:
                    if os.path.exists(tmp_img_path):
                        os.remove(tmp_img_path)

            except Exception as e:
                logger.warning(f"Error processing embedded image {img_data['idx']} on page {page_num + 1}: {e}")
                continue

    except Exception as e:
        logger.error(f"Error extracting embedded images from page {page_num + 1}: {e}")

    return ocr_results, total_ocr_count


# --- WORKER FUNCTION ---
def worker_process_pdf_from_bytes(
    pdf_bytes: bytes,
    source_url: str,
    gemini_api_key: Optional[str],
    start_page: int = 0,
    end_page: Optional[int] = None,
    memory_limit_mb: int = 1024,
    auto_ocr: bool = True
) -> Optional[str]:
    """
    Worker function for PDF processing that can be run in a separate process.
    Handles strict memory management and CPU-intensive OCR tasks.
    """
    # Setup logging in worker process
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    logger = logging.getLogger(f"PDFWorker-{os.getpid()}")
    
    pdf_size_mb = len(pdf_bytes) / (1024 * 1024)
    logger.info(f"Processing PDF from bytes: {source_url} ({pdf_size_mb:.1f}MB)")
    
    # Setup Gemini client locally in worker
    gemini_model = setup_gemini_client(gemini_api_key)
    
    markdown_parts = []
    total_images_processed = 0
    scanned_pages_ocrd = 0
    
    ocr_available = gemini_model is not None
    
    # Check if OCR is enabled and auto-detect settings
    # If auto_ocr is True, we use smart detection
    # If auto_ocr is False, we skip all OCR
    # Legacy ocr_enabled parameter is handled by caller (iterative_processor) and passed as auto_ocr
    
    skip_ocr_for_memory = pdf_size_mb > 50
    if skip_ocr_for_memory and ocr_available:
        logger.warning(f"PDF is large ({pdf_size_mb:.1f}MB), limiting OCR to save memory")
        
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        total_doc_pages = len(doc)
        
        # Determine actual end page
        real_end_page = total_doc_pages if end_page is None else min(end_page, total_doc_pages)
        real_start_page = max(0, start_page)
        
        if real_start_page >= real_end_page:
            logger.info(f"No pages to process in range {real_start_page}-{real_end_page} (Total: {total_doc_pages})")
            return ""

        logger.info(f"Processing PDF pages {real_start_page}-{real_end_page} (Total: {total_doc_pages})")
            
        for page_num in range(real_start_page, real_end_page):
            page_content_parts = []
            
            # Memory check
            if page_num > 0 and page_num % 10 == 0:
                if not check_memory_usage(memory_limit_mb):
                    logger.warning(f"Memory limit reached at page {page_num + 1}, skipping remaining OCR")
                    skip_ocr_for_memory = True
                if pdf_size_mb > 20:
                    cleanup_memory()
            
            page = doc.load_page(page_num)
            text_content = page.get_text("text").strip()
            
            # Use module-level function
            image_list = page.get_images(full=True)
            has_images = has_significant_images(doc, image_list)
            
            if text_content:
                page_content_parts.append(text_content)
                
                # Auto-OCR Charts
                if auto_ocr and ocr_available and has_images and not skip_ocr_for_memory:
                    embedded_images, count = extract_and_ocr_embedded_images_memory(
                        doc, page_num, source_url, gemini_model,
                        min_size=200, max_images_per_page=5,
                        total_ocr_count=total_images_processed
                    )
                    total_images_processed = count
                    if embedded_images:
                        page_content_parts.append("\n### Embedded Graphics (Auto-Extracted)\n")
                        page_content_parts.extend(embedded_images)
                        
            elif auto_ocr and ocr_available and not skip_ocr_for_memory:
                # Scanned Page OCR
                scanned_pages_ocrd += 1
                try:
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                        tmp_img.write(pix.tobytes("png"))
                        tmp_img_path = tmp_img.name
                        
                    try:
                        ocr_content = ocr_image_to_markdown(tmp_img_path, gemini_model)
                        if ocr_content:
                            page_content_parts.append(ocr_content)
                    finally:
                        if os.path.exists(tmp_img_path):
                            os.remove(tmp_img_path)
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num}: {e}")
                    page_content_parts.append(f"[Page {page_num + 1}: Scanned content, OCR failed]")

            if page_content_parts:
                page_markdown = "\n".join(page_content_parts)
                markdown_parts.append(f"## Page {page_num + 1}\n\n{page_markdown}\n")
                
        doc.close()
        cleanup_memory()
        
        final_markdown = f"# PDF: {source_url}\n\n" + "\n---\n\n".join(markdown_parts)
        return final_markdown
        
    except Exception as e:
        logger.error(f"Error processing PDF from bytes ({source_url}): {e}")
        cleanup_memory()
        return None


# --- Main Class ---

class PDFProcessor:
    """Handles PDF and document processing with OCR and AI."""
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.gemini_model = setup_gemini_client(os.getenv("GEMINI_API_KEY"))

    def process_pdf(self, pdf_path: Path, output_dir: Path, temp_image_dir: Path) -> Optional[Path]:
        """Process a single PDF file (synchronous, file-based)."""
        logger.info(f"Processing PDF: {pdf_path}")
        pdf_stem = pdf_path.stem
        markdown_parts = []
        temp_files_to_clean = []

        try:
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page_content_parts = []
                logger.info(f"Processing page {page_num + 1}/{len(doc)} of {pdf_path.name}")
                
                # Image for OCR
                page_image_filename = f"{pdf_stem}_page_{page_num + 1}_full.png"
                page_image_path = temp_image_dir / page_image_filename
                
                page = doc.load_page(page_num)
                pix = page.get_pixmap()
                pix.save(str(page_image_path))
                temp_files_to_clean.append(page_image_path)

                # OCR
                page_markdown = ocr_image_to_markdown(str(page_image_path), self.gemini_model, is_embedded_graphic=False)
                if page_markdown:
                    page_content_parts.append(page_markdown)

                # Embedded images
                embedded_image_paths = self.extract_embedded_images_from_page(doc, page_num, pdf_stem, temp_image_dir)
                
                if embedded_image_paths:
                    page_content_parts.append("\n\n--- Embedded Graphics Content ---\n")
                    for emb_img_path in embedded_image_paths:
                        temp_files_to_clean.append(emb_img_path)
                        embedded_image_markdown = ocr_image_to_markdown(str(emb_img_path), self.gemini_model, is_embedded_graphic=True)
                        
                        if embedded_image_markdown:
                            page_content_parts.append(f"\n**Content from embedded graphic {emb_img_path.name}:**\n")
                            page_content_parts.append(embedded_image_markdown)
                            page_content_parts.append("\n")
                    page_content_parts.append("--- End Embedded Graphics Content ---\n\n")
                
                markdown_parts.append("".join(page_content_parts))
                markdown_parts.append(f"\n\n{'='*20} Page {page_num + 1} {'='*20}\n\n")

            doc.close()
            final_markdown_content = "".join(markdown_parts)
            
            output_dir.mkdir(exist_ok=True)
            output_md_path = output_dir / f"{pdf_stem}.md"
            
            with open(output_md_path, "w", encoding="utf-8") as f:
                f.write(final_markdown_content)
            
            return output_md_path

        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {e}")
            return None
        finally:
            for temp_file in temp_files_to_clean:
                try:
                    if temp_file.exists():
                        os.remove(temp_file)
                except Exception:
                    pass

    def extract_embedded_images_from_page(self, doc, page_num: int, pdf_stem: str, output_dir: Path) -> List[Path]:
        """Extract embedded images from a PDF page (file-based)."""
        extracted_image_paths = []
        page = doc.load_page(page_num)
        image_list = page.get_images(full=True)

        for img_idx, img_info in enumerate(image_list):
            xref = img_info[0]
            try:
                base_image = doc.extract_image(xref)
                if not base_image:
                    continue
                    
                image_bytes = base_image["image"]
                ext = base_image["ext"]
                
                safe_pdf_stem = "".join(c if c.isalnum() or c in ('_','-') else '_' for c in pdf_stem)
                image_filename = f"{safe_pdf_stem}_page_{page_num + 1}_embedded_img_{img_idx}.{ext}"
                image_path = output_dir / image_filename
                
                with open(image_path, "wb") as img_file:
                    img_file.write(image_bytes)
                extracted_image_paths.append(image_path)
                
            except Exception as e:
                logger.error(f"Error extracting image xref {xref}: {e}")
                
        return extracted_image_paths

    def process_pdf_from_bytes(self, pdf_bytes: bytes, source_url: str,
                                ocr_enabled: bool = None, 
                                start_page: int = 0,
                                end_page: Optional[int] = None,
                                memory_limit_mb: int = 1024,
                                auto_ocr: bool = True) -> Optional[str]:
        """
        Process PDF from bytes in memory.
        Delegates to the worker function (synchronously here, but design allows external async use).
        """
        return worker_process_pdf_from_bytes(
            pdf_bytes=pdf_bytes,
            source_url=source_url,
            gemini_api_key=os.getenv("GEMINI_API_KEY"),
            start_page=start_page,
            end_page=end_page,
            memory_limit_mb=memory_limit_mb,
            auto_ocr=True if ocr_enabled is None else auto_ocr
        )

    def process_word_document(self, docx_path: Path, output_dir: Path) -> Optional[Path]:
        """Process a Word document."""
        logger.info(f"Processing Word document: {docx_path}")
        try:
            try:
                from docx import Document
                doc = Document(docx_path)
                markdown_content = f"# {docx_path.stem}\n\n"
                
                for para in doc.paragraphs:
                    if para.style.name.startswith('Heading'):
                        level = int(para.style.name[-1]) if para.style.name[-1].isdigit() else 1
                        markdown_content += f"{'#' * level} {para.text}\n\n"
                    else:
                        if para.text.strip():
                            markdown_content += f"{para.text}\n\n"
                
                for table in doc.tables:
                    markdown_table = "| "
                    if len(table.rows) > 0:
                        for cell in table.rows[0].cells:
                            markdown_table += cell.text + " | "
                        markdown_table += "\n| "
                        for _ in table.rows[0].cells:
                            markdown_table += "--- | "
                        markdown_table += "\n"
                        for row_idx, row in enumerate(table.rows):
                            if row_idx == 0: continue
                            markdown_table += "| "
                            for cell in row.cells:
                                markdown_table += cell.text + " | "
                            markdown_table += "\n"
                        markdown_content += markdown_table + "\n\n"
                
                output_dir.mkdir(exist_ok=True)
                output_md_path = output_dir / f"{docx_path.stem}.md"
                with open(output_md_path, "w", encoding="utf-8") as f:
                    f.write(markdown_content)
                return output_md_path
                
            except ImportError:
                logger.warning("python-docx not available, using Gemini for Word document")
                return self._process_with_gemini(docx_path, output_dir, "Word document")
                
        except Exception as e:
            logger.error(f"Error processing Word document {docx_path}: {e}")
            return None

    def process_excel_file(self, excel_path: Path, output_dir: Path) -> Optional[Path]:
        """Process an Excel file."""
        logger.info(f"Processing Excel file: {excel_path}")
        try:
            excel_file = pd.ExcelFile(excel_path)
            sheet_names = excel_file.sheet_names
            markdown_content = f"# {excel_path.stem}\n\n"
            for sheet_name in sheet_names:
                df = pd.read_excel(excel_path, sheet_name=sheet_name)
                markdown_content += f"## Sheet: {sheet_name}\n\n"
                markdown_content += df.to_markdown(index=False) + "\n\n"
            
            output_dir.mkdir(exist_ok=True)
            output_md_path = output_dir / f"{excel_path.stem}.md"
            with open(output_md_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)
            return output_md_path
        except Exception as e:
            logger.error(f"Error processing Excel file {excel_path}: {e}")
            return None

    def process_csv_file(self, csv_path: Path, output_dir: Path) -> Optional[Path]:
        """Process a CSV file."""
        logger.info(f"Processing CSV file: {csv_path}")
        try:
            df = pd.read_csv(csv_path)
            markdown_content = f"# {csv_path.stem}\n\n"
            markdown_content += df.to_markdown(index=False)
            output_dir.mkdir(exist_ok=True)
            output_md_path = output_dir / f"{csv_path.stem}.md"
            with open(output_md_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)
            return output_md_path
        except Exception as e:
            logger.error(f"Error processing CSV file {csv_path}: {e}")
            return None

    def _process_with_gemini(self, file_path: Path, output_dir: Path, file_type: str) -> Optional[Path]:
        """Process a file using Gemini AI as fallback."""
        if not self.gemini_model:
            return None
        try:
            response = self.gemini_model.generate_content([
                f"Convert this {file_type} to markdown. Use appropriate markdown syntax for headings, lists, tables, and formatting. Preserve original structure and layout. Do NOT wrap the content in markdown code blocks or use triple backticks.",
                genai.upload_file(str(file_path))
            ])
            cleaned_text = response.text.strip()
            if cleaned_text.startswith("```markdown"): cleaned_text = cleaned_text[11:].strip()
            elif cleaned_text.startswith("```"): cleaned_text = cleaned_text[3:].strip()
            if cleaned_text.endswith("```"): cleaned_text = cleaned_text[:-3].strip()
            
            output_dir.mkdir(exist_ok=True)
            output_md_path = output_dir / f"{file_path.stem}.md"
            with open(output_md_path, "w", encoding="utf-8") as f:
                f.write(cleaned_text)
            return output_md_path
        except Exception as e:
            logger.error(f"Error processing {file_type} with Gemini: {e}")
            return None

    def process_documents(self, input_dir: str, output_dir: str, temp_dir: str) -> Dict[str, Any]:
        """Process all documents in a directory."""
        logger.info(f"Processing documents from {input_dir} to {output_dir}")
        input_path = Path(input_dir)
        files_to_process = []
        for ext in ['.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv']:
            files_to_process.extend(input_path.rglob(f"*{ext}"))
        
        processed_files = []
        failed_files = []
        
        def process_single_document(file_path: Path):
            try:
                ext = file_path.suffix.lower()
                if ext == '.pdf': result = self.process_pdf(file_path, Path(output_dir), Path(temp_dir))
                elif ext in ['.docx', '.doc']: result = self.process_word_document(file_path, Path(output_dir))
                elif ext in ['.xlsx', '.xls']: result = self.process_excel_file(file_path, Path(output_dir))
                elif ext == '.csv': result = self.process_csv_file(file_path, Path(output_dir))
                else: return False
                
                if result: processed_files.append({"source": str(file_path), "output": str(result)})
                else: failed_files.append(str(file_path))
            except Exception: failed_files.append(str(file_path))

        with ThreadPoolExecutor(max_workers=self.config.pdf_concurrency) as executor:
            executor.map(process_single_document, files_to_process)
            
        return {
            "total_files": len(files_to_process),
            "processed_files": len(processed_files),
            "failed_files": len(failed_files),
            "results": processed_files
        }
